from flask import Flask, render_template, request, send_file
import os, subprocess, shutil
from pdf2image import convert_from_path
import pytesseract
import camelot
from docx import Document
from PyPDF2 import PdfReader

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def ocr_pdf(input_path, output_path):
    """
    Run OCRmyPDF to add searchable text layer to scanned PDFs.
    Falls back gracefully if OCRmyPDF is not installed (Windows dev mode).
    """
    if shutil.which("ocrmypdf") is None:
        print("⚠️ OCRmyPDF not found. Skipping OCR step.")
        return input_path  # fallback: return original PDF

    try:
        subprocess.run(
            ["ocrmypdf", "--skip-text", "--force-ocr", input_path, output_path],
            check=True
        )
        return output_path
    except Exception as e:
        print(f"⚠️ OCRmyPDF failed: {e}")
        return input_path


@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        file = request.files["pdf"]
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)

        # Step 1: Check if PDF has embedded text
        reader = PdfReader(filepath)
        has_text = any(page.extract_text() for page in reader.pages)

        processed_pdf = filepath
        if not has_text:  # scanned PDF → run OCR if possible
            processed_pdf = filepath.replace(".pdf", "_ocr.pdf")
            processed_pdf = ocr_pdf(filepath, processed_pdf)

        # Step 2: Extract text
        text_pages = []
        pages = convert_from_path(processed_pdf, dpi=200)
        for page in pages:
            text_pages.append(pytesseract.image_to_string(page))

        # Step 3: Extract tables
        try:
            tables = camelot.read_pdf(processed_pdf, pages="all")
            tables_preview = [t.df.to_html(classes="table-preview") for t in tables]
        except:
            tables_preview = []

        return render_template(
            "preview.html",
            text_pages=text_pages,
            tables=tables_preview,
            filename=file.filename
        )

    return render_template("upload.html")


@app.route("/confirm/<filename>", methods=["POST"])
def confirm(filename):
    filepath = os.path.join(UPLOAD_FOLDER, filename)

    # Check if scanned
    reader = PdfReader(filepath)
    has_text = any(page.extract_text() for page in reader.pages)

    processed_pdf = filepath
    if not has_text:
        processed_pdf = filepath.replace(".pdf", "_ocr.pdf")
        processed_pdf = ocr_pdf(filepath, processed_pdf)

    # Build Word doc
    doc = Document()
    pages = convert_from_path(processed_pdf, dpi=200)
    for page in pages:
        doc.add_paragraph(pytesseract.image_to_string(page))

    try:
        tables = camelot.read_pdf(processed_pdf, pages="all")
        for t in tables:
            table = doc.add_table(rows=0, cols=len(t.df.columns))
            for row in t.df.values.tolist():
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
    except:
        pass

    output_path = os.path.join(OUTPUT_FOLDER, filename.replace(".pdf", ".docx"))
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
